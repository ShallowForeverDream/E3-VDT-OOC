
from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "report" / "FINAL_REPORT_DRAFT.md"
OUT = ROOT / "docs" / "report" / "E3-VDT-OOC-结课报告初稿.docx"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C4D2", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_code
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("E3-VDT-OOC 结课项目  |  第 ")
    set_run_font(r, size=9, color="666666")
    add_field(p, "PAGE")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color="666666")


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_markdown_paragraph(doc: Document, text: str):
    if not text.strip():
        return
    stripped = text.strip()
    if stripped.startswith(">"):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(stripped.lstrip("> "))
        set_run_font(run, size=11, bold=True, color="1F3A5F")
        return
    if re.match(r"^- ", stripped):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(clean_inline(stripped[2:]))
        set_run_font(run)
        return
    if re.match(r"^\d+\. ", stripped):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(clean_inline(re.sub(r"^\d+\.\s*", "", stripped)))
        set_run_font(run)
        return
    p = doc.add_paragraph()
    # simple bold/code parser
    pos = 0
    pattern = re.compile(r"(\*\*([^*]+)\*\*|`([^`]+)`)")
    for m in pattern.finditer(stripped):
        if m.start() > pos:
            r = p.add_run(stripped[pos:m.start()])
            set_run_font(r)
        if m.group(2):
            r = p.add_run(m.group(2))
            set_run_font(r, bold=True)
        else:
            r = p.add_run(m.group(3))
            set_run_font(r, name="Consolas", east_asia="Microsoft YaHei", size=10, color="1F4D78")
        pos = m.end()
    if pos < len(stripped):
        r = p.add_run(stripped[pos:])
        set_run_font(r)


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_table(lines: List[str]):
    rows = []
    for line in lines:
        cells = [clean_inline(c.strip()) for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and is_table_separator(lines[1]):
        rows = [rows[0]] + rows[2:]
    return rows


def add_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    # simple equal widths, explicit geometry via cell width
    width = Inches(6.5 / max(cols, 1))
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 or len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9 if cols >= 5 else 10, bold=(i == 0), color="0B2545" if i == 0 else "000000")
    doc.add_paragraph()


def build():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)

    # title block
    title = "基于事件语境与证据约束的跨域图文内容挪用检测系统"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color="0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("E3-VDT-OOC 期末项目报告初稿")
    set_run_font(r, size=12, color="555555")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("分类准确率不降，结构化归因增强")
    set_run_font(r, size=11, bold=True, color="2E74B5")
    doc.add_paragraph()

    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            for idx, line in enumerate(code_buf):
                if idx:
                    p.add_run("\n")
                r = p.add_run(line)
                set_run_font(r, name="Consolas", east_asia="Microsoft YaHei", size=9.5, color="1F4D78")
            code_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, parse_table(table_buf))
            table_buf = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
                code_buf = []
            continue
        if in_code:
            code_buf.append(line)
            continue
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buf.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            # skip duplicate source doc title
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:].strip()), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:].strip()), level=2)
        elif line.startswith("#### "):
            doc.add_heading(clean_inline(line[5:].strip()), level=3)
        else:
            add_markdown_paragraph(doc, line)
    flush_table(); flush_code()

    doc.core_properties.title = title
    doc.core_properties.subject = "内容安全课程期末项目"
    doc.core_properties.author = "E3-VDT-OOC Team"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

