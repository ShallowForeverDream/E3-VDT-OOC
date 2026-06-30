from __future__ import annotations

import csv
import json
import math
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPO = ROOT / "E3-VDT-OOC"
SUBMIT_ROOT = ROOT / "SUBMIT" / "10-VDT-CF-Attr跨域图文内容挪用检测与错配归因-朱俊好"
REPORT_DIR = SUBMIT_ROOT / "1_实验报告"
ASSET_DIR = REPORT_DIR / "paper_assets"

TITLE = "基于可控反事实训练的跨域图文内容挪用检测与错配归因方法"
EN_TITLE = "Counterfactual Attribution for Cross-Domain Out-of-Context Image-Text Misinformation Detection"
MEMBERS = [
    ("费思岳", "2023302181043"),
    ("陈宇欣", "2023302181266"),
    ("朱俊好", "2023335550135"),
    ("郭远重", "2023302181159"),
]


def load_json(rel: str) -> dict:
    return json.loads((REPO / rel).read_text(encoding="utf-8"))


def fmt(x: float) -> str:
    return f"{x:.4f}"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt, fill=(0, 0, 0)):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total_h = sum(heights) + (len(lines) - 1) * 6
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 6


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill=(20, 70, 120)):
    draw.line([start, end], fill=fill, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 14
    pts = [
        end,
        (end[0] - head * math.cos(angle - math.pi / 6), end[1] - head * math.sin(angle - math.pi / 6)),
        (end[0] - head * math.cos(angle + math.pi / 6), end[1] - head * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(pts, fill=fill)


def save_pipeline(path: Path):
    img = Image.new("RGB", (1800, 660), "white")
    d = ImageDraw.Draw(img)
    box_font = font(28)
    small_font = font(22)
    boxes = [
        ((80, 120, 360, 270), "输入\n图片 + 新闻文本"),
        ((470, 120, 750, 270), "VDTAdapter\nOOC / Non-OOC"),
        ((860, 120, 1170, 270), "归因门控\n仅 OOC 进入"),
        ((1280, 120, 1660, 270), "错配归因\n类型 + 冲突字段"),
        ((470, 410, 780, 550), "COVE-lite\n上下文构造与上限评估"),
        ((930, 410, 1300, 550), "可控反事实训练\n单字段 + different-event"),
    ]
    for box, label in boxes:
        d.rounded_rectangle(box, radius=20, fill=(245, 248, 252), outline=(20, 70, 120), width=3)
        center_text(d, box, label, box_font)
    arrow(d, (360, 195), (470, 195))
    arrow(d, (750, 195), (860, 195))
    arrow(d, (1170, 195), (1280, 195))
    arrow(d, (625, 410), (930, 480))
    arrow(d, (1115, 410), (1020, 270))
    d.text((88, 315), "推理阶段不输入 true image context；Non-OOC 与 Uncertain 不强行归因。", font=small_font, fill=(60, 60, 60))
    d.text((470, 580), "训练阶段使用上下文与反事实构造监督信号，真实 OOC 人工集只用于评估。", font=small_font, fill=(60, 60, 60))
    img.save(path)


def chart_bar(path: Path, title: str, labels: list[str], values: list[float], ylabel: str, ymax: float | None = None):
    w, h = 1500, 700
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    axis_f, label_f = font(23), font(22)
    left, top, right, bottom = 140, 75, 1420, 555
    d.line([(left, top), (left, bottom), (right, bottom)], fill=(0, 0, 0), width=3)
    ymax = ymax or max(values) * 1.15
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        d.line([(left, y), (right, y)], fill=(225, 225, 225), width=1)
        tick = f"{ymax * i / 5:.2f}" if ymax <= 1 else f"{ymax * i / 5:.0f}"
        d.text((45, y - 12), tick, font=axis_f, fill=(0, 0, 0))
    n = len(values)
    gap = 42
    bw = (right - left - gap * (n + 1)) / n
    colors = [(31, 119, 180), (76, 120, 168), (114, 183, 178), (245, 133, 24), (228, 87, 86), (90, 90, 90)]
    for i, (lab, val) in enumerate(zip(labels, values)):
        x1 = left + gap + i * (bw + gap)
        x2 = x1 + bw
        y1 = bottom - (bottom - top) * val / ymax
        d.rectangle((x1, y1, x2, bottom), fill=colors[i % len(colors)])
        txt = f"{val:.3f}" if ymax <= 1 else str(int(val))
        d.text((x1 + bw / 2 - text_size(d, txt, label_f)[0] / 2, y1 - 34), txt, font=label_f, fill=(0, 0, 0))
        for j, part in enumerate(lab.split("\n")):
            d.text((x1 + bw / 2 - text_size(d, part, axis_f)[0] / 2, bottom + 22 + j * 28), part, font=axis_f, fill=(0, 0, 0))
    d.text((20, top + 170), ylabel, font=axis_f, fill=(0, 0, 0))
    img.save(path)


def chart_lines(path: Path, title: str, xs: list[int], series: list[tuple[str, list[float], tuple[int, int, int]]]):
    w, h = 1500, 700
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    axis_f, label_f = font(23), font(22)
    left, top, right, bottom = 140, 75, 1350, 555
    ymax = 0.65
    d.line([(left, top), (left, bottom), (right, bottom)], fill=(0, 0, 0), width=3)
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        d.line([(left, y), (right, y)], fill=(225, 225, 225), width=1)
        d.text((55, y - 12), f"{ymax * i / 5:.2f}", font=axis_f, fill=(0, 0, 0))
    x_positions = [left + (right - left) * i / (len(xs) - 1) for i in range(len(xs))]
    for name, vals, color in series:
        pts = [(x_positions[i], bottom - (bottom - top) * vals[i] / ymax) for i in range(len(vals))]
        d.line(pts, fill=color, width=5)
        for (x, y), val in zip(pts, vals):
            d.ellipse((x - 9, y - 9, x + 9, y + 9), fill=color)
            d.text((x - 28, y - 38), f"{val:.3f}", font=label_f, fill=(0, 0, 0))
    for x, lab in zip(x_positions, xs):
        d.text((x - 25, bottom + 24), str(lab), font=axis_f, fill=(0, 0, 0))
    d.text((610, bottom + 70), "MaxPerType", font=axis_f, fill=(0, 0, 0))
    for idx, (name, _vals, color) in enumerate(series):
        y = 155 + idx * 42
        d.line([(1160, y), (1210, y)], fill=color, width=5)
        d.text((1225, y - 14), name, font=axis_f, fill=(0, 0, 0))
    img.save(path)


def chart_grouped(path: Path, title: str, groups: list[str], metrics: list[tuple[str, list[float], tuple[int, int, int]]]):
    w, h = 1500, 700
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    axis_f, label_f = font(23), font(22)
    left, top, right, bottom = 170, 75, 1350, 555
    ymax = 0.55
    d.line([(left, top), (left, bottom), (right, bottom)], fill=(0, 0, 0), width=3)
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        d.line([(left, y), (right, y)], fill=(225, 225, 225), width=1)
        d.text((70, y - 12), f"{ymax * i / 5:.2f}", font=axis_f, fill=(0, 0, 0))
    group_w = (right - left) / len(groups)
    bar_w = 90
    for gi, g in enumerate(groups):
        center = left + group_w * (gi + 0.5)
        d.text((center - text_size(d, g, axis_f)[0] / 2, bottom + 26), g, font=axis_f, fill=(0, 0, 0))
        for mi, (name, vals, color) in enumerate(metrics):
            x1 = center + (mi - (len(metrics) - 1) / 2) * (bar_w + 20) - bar_w / 2
            val = vals[gi]
            y1 = bottom - (bottom - top) * val / ymax
            d.rectangle((x1, y1, x1 + bar_w, bottom), fill=color)
            d.text((x1 + bar_w / 2 - 28, y1 - 34), f"{val:.3f}", font=label_f, fill=(0, 0, 0))
    for idx, (name, _vals, color) in enumerate(metrics):
        y = 155 + idx * 42
        d.rectangle((1145, y - 12, 1185, y + 12), fill=color)
        d.text((1200, y - 14), name, font=axis_f, fill=(0, 0, 0))
    img.save(path)


def crop_white_margin(path: Path, pad: int = 18):
    img = Image.open(path).convert("RGB")
    w, h = img.size
    pixels = img.load()
    xs, ys = [], []
    for y in range(h):
        for x in range(w):
            r, g, b = pixels[x, y]
            if r < 250 or g < 250 or b < 250:
                xs.append(x)
                ys.append(y)
    if not xs:
        return
    box = (
        max(0, min(xs) - pad),
        max(0, min(ys) - pad),
        min(w, max(xs) + pad + 1),
        min(h, max(ys) + pad + 1),
    )
    img.crop(box).save(path)


def build_assets(metrics: dict) -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "pipeline": ASSET_DIR / "fig1_pipeline.png",
        "scaling": ASSET_DIR / "fig2_scaling.png",
        "manual": ASSET_DIR / "fig3_manual_distribution.png",
        "real": ASSET_DIR / "fig4_real_eval.png",
    }
    save_pipeline(paths["pipeline"])
    chart_lines(
        paths["scaling"],
        "no-true-context 归因训练规模曲线",
        [80, 200, 1000],
        [
            ("Type Acc", [0.2745, 0.4266, 0.5275], (31, 119, 180)),
            ("Field Micro-F1", [0.3564, 0.5195, 0.5719], (114, 183, 178)),
        ],
    )
    manual = metrics["manual"]
    chart_bar(
        paths["manual"],
        "真实 OOC 人工 100 条错配类型分布",
        ["different-\nevent", "entity", "location", "context\nomission", "event-\ntype"],
        [
            manual["type_distribution"]["different-event mismatch"],
            manual["type_distribution"]["entity mismatch"],
            manual["type_distribution"]["location mismatch"],
            manual["type_distribution"]["context omission"],
            manual["type_distribution"]["event-type mismatch"],
        ],
        "样本数",
        ymax=90,
    )
    real_eval = metrics["real_eval"]["models"]
    chart_grouped(
        paths["real"],
        "真实 OOC 100 条归因评估",
        ["5way_1000", "plus2000"],
        [
            (
                "Type Acc",
                [
                    real_eval["no_true_context_attr_5way_1000"]["mismatch_type_accuracy"],
                    real_eval["no_true_context_attr_5way_plus2000"]["mismatch_type_accuracy"],
                ],
                (76, 120, 168),
            ),
            (
                "Field Micro-F1",
                [
                    real_eval["no_true_context_attr_5way_1000"]["conflict_field_micro_f1"],
                    real_eval["no_true_context_attr_5way_plus2000"]["conflict_field_micro_f1"],
                ],
                (114, 183, 178),
            ),
        ],
    )
    for p in paths.values():
        crop_white_margin(p)
    return paths


@dataclass
class Table:
    caption: str
    header: list[str]
    rows: list[list[str]]
    note: str | None = None
    widths: list[float] | None = None


@dataclass
class Figure:
    caption: str
    path: Path
    note: str


def content(metrics: dict, figs: dict[str, Path]):
    plus = metrics["plus"]
    five = metrics["five"]
    manual = metrics["manual"]
    real_eval = metrics["real_eval"]["models"]
    ctx = metrics["ctx"]
    event = metrics["event"]

    abstract = [
        "真实新闻图片与错误文本语境重新配对后，可以在不修改像素的条件下改变事件含义。现有跨域 OOC 检测模型多输出 OOC/Non-OOC 标签，能够判断图文是否错配，却难以定位主体、地点、时间、事件类型或行为关系中的冲突字段。本文在 VDT 跨域检测基线上研究 no-true-context 条件下的错配归因问题：推理阶段只输入图片和当前文本，不输入图片原始新闻上下文。",
        "本文采用主分类与归因解耦的协议。VDT 负责 OOC 主分类，归因头只在 OOC 条件下预测 mismatch type 与 conflict fields；Non-OOC 或证据不足样本不生成强归因。由于 NewsCLIPpings 仅提供二分类标签，本文利用 VisualNews 原始上下文构造 COVE-lite oracle 链路，用于样本构造和上限分析；同时从 Non-OOC 样本生成 entity、location、time 单字段反事实，并从原始 OOC 中筛选低相似、低重叠的 different-event 样本。",
        "VDT strict BLIP-2/GaussianBlur 复现在 bbc,guardian 目标域取得 F1 0.7353、Acc 0.7383、AUC 0.7398，在 usa_today,washington_post 目标域 batch size 64 设置下取得 F1 0.8032、Acc 0.8032、AUC 0.8028。no-true-context 归因头在 MaxPerType=80/200/1000 三组反事实训练中，Type Acc 从 0.2745 升至 0.5275，Field Micro-F1 从 0.3564 升至 0.5719。加入 2000 条 filtered different-event 候选后，合成测试集 Field Micro-F1 为 0.6876；真实 OOC 100 条上 Field Micro-F1 从 0.3276 升至 0.4781，Exact Match 仍为 0.0300。人工标注显示真实 OOC 中 different-event 占 85%，且常与多个字段共现，说明归因误差主要来自多字段事件结构恢复，而不是单字段替换模式本身。",
    ]

    sections = [
        (
            "1 引言",
            [
                "Out-of-Context 图文内容挪用将真实图片与错误文本组合到同一叙事中。旧图新用、同主体不同事件、同主题跨地点配图和灾害现场错误引用都属于该类风险。图像和文本可以分别真实，误导来自二者绑定到不同事件。高语义相似的 hard negative 使相似度匹配不足以承担事实一致性判断。",
                "内容安全审核需要两个层次的输出：图文是否属于同一事件，以及不一致时冲突发生在哪些事件字段。前者对应 OOC 二分类，后者对应可审计的错配归因。本文围绕三个可检验问题展开：VDT strict 设置在本地硬件约束下能否复现跨域检测结果；no-true-context 推理中，归因头能否从可控反事实中学习字段冲突；真实 OOC 的错配结构是否仍能由单字段反事实解释。",
                "本文的主要贡献有三点。第一，复现 VDT strict BLIP-2/GaussianBlur 两组跨域实验，并记录 batch size 受显存约束后的可运行设置。第二，构建 VDT-CF-Attr 后置归因协议，使 mismatch type 和 conflict fields 在 OOC 门控下输出，不覆盖主分类。第三，构造单字段反事实与 filtered different-event 训练集，并用 100 条真实 OOC 人工样本检验合成分布到真实分布的变化。",
            ],
            [],
            [],
        ),
        (
            "2 相关工作",
            [
                "NewsCLIPpings 基于 VisualNews 重新配对真实新闻图片和文本，形成 OOC 检测基准。该数据集刻画“真实图像 + 错误文本语境”的误导传播，但标签粒度停留在 OOC/Non-OOC。VisualNews 保存图片原始新闻上下文，因此可以为错配字段构造提供参照。",
                "跨域 OOC 检测关注新闻机构、话题和写作风格变化下的泛化。VDT 通过变分域不变表示和测试时训练缓解域偏移，适合作为主分类基线。解释型 OOC 研究进一步引入外部证据或图像原始上下文，代表工作包括 SNIFFER 和 COVE。这类方法表明，语境一致性需要比较事件主体、地点、时间和关系，而不是只比较图文相似度。",
                "相似度 shortcut 仍可能在 OOC benchmark 中形成强信号，尤其在负样本构造规则固定时。本文将检测分数、错配字段和真实事件一致性分开评估：VDT 用于主分类，COVE-lite 用于构造和 oracle 分析，最终归因头只使用 image+caption 输入。",
            ],
            [
                Table(
                    "表1  数据来源与评价对象",
                    ["数据源", "原始标签", "使用环节"],
                    [
                        ["NewsCLIPpings", "OOC/Non-OOC", "VDT 复现、反事实构造、真实 OOC 抽样"],
                        ["VisualNews", "图像原始上下文", "COVE-lite oracle、人工标注参考"],
                        ["人工 OOC 100 条", "mismatch type；conflict fields", "真实归因评估"],
                    ],
                    "注：NewsCLIPpings 原始标签不包含错配字段，本文只把人工标注作为真实归因评估依据。",
                    [1.7, 2.0, 2.8],
                )
            ],
            [],
        ),
        (
            "3 方法",
            [
                "给定图片 I 与当前新闻文本 T，VDTAdapter 输出主分类 y。若 y 为 OOC，归因头 h_attr(I,T) 输出错配类型 a 和冲突字段集合 F；若 y 为 Non-OOC，系统输出 benign；若 y 为 Uncertain，系统输出 evidence insufficient。该门控结构防止归因模块覆盖主分类结果。",
                "COVE-lite 链路通过 image_id 找到 VisualNews 原始上下文 C，并比较 T 与 C 的事件字段。该链路不进入最终推理输入，用于构造反事实样本、分析 oracle 上限和辅助人工标注。事件字段包括 entity、location、time、event_type 和 relation。",
                "可控反事实样本从 Non-OOC 图文对生成。保持图片不变，只替换文本中的实体、地点或年份字段，分别标记为 entity mismatch、location mismatch 和 temporal mismatch。正常样本标为 benign illustrative image。对原始 OOC 样本，只选择低文本相似度、低 token 重叠且不与人工评估集重合的样本作为 different-event 候选。",
                "no-true-context 归因头的输入特征由图文相似度、字段 prompt grounding、文本字段存在性和 VDT score 组成。训练候选包括规则基线、logistic regression 与多层感知机。默认模型由测试集 Type Acc、Field Micro-F1 和 Exact Match 综合选择。",
            ],
            [
                Table(
                    "表2  归因标签与训练来源",
                    ["标签", "样本来源", "监督性质"],
                    [
                        ["benign", "Non-OOC 原样本", "原始匹配样本"],
                        ["entity", "实体替换", "单字段反事实"],
                        ["location", "地点替换", "单字段反事实"],
                        ["temporal", "年份替换", "单字段反事实"],
                        ["different-event", "筛选原始 OOC", "弱监督候选"],
                    ],
                    "注：different-event 样本在 plus2000 设置中排除了人工 100 条评估集。",
                    [1.6, 2.4, 2.5],
                )
            ],
            [
                Figure(
                    "图1  VDT-CF-Attr 的训练与推理结构",
                    figs["pipeline"],
                    "图1给出主分类、归因门控和反事实训练之间的关系。推理阶段不读取 true image context；COVE-lite 只参与训练构造、oracle 分析和人工标注参考。",
                )
            ],
        ),
        (
            "4 实验设置",
            [
                "实验分为四组。第一组复现 VDT strict BLIP-2/GaussianBlur 跨域检测结果。第二组构造 COVE-lite context pairs 并检查字段抽取覆盖率。第三组在 no-true-context 设置下比较反事实训练规模。第四组在真实 OOC 人工 100 条上评估归因泛化。",
                "主要指标包括 Type Acc、Field Micro-F1 和 Exact Match。Type Acc 评价主错配类型是否正确；Field Micro-F1 评价冲突字段集合；Exact Match 要求类型和字段同时匹配。真实 OOC 的 Exact Match 更严格，因为 different-event 样本常同时涉及多个字段。本文列出点估计和误差分布，不把单次课程实验结果写成统计显著性结论。",
            ],
            [
                Table(
                    "表3  VDT strict baseline 复现结果",
                    ["目标域", "Batch", "F1", "Acc", "AUC", "记录"],
                    [
                        ["bbc, guardian", "128", "0.7353", "0.7383", "0.7398", "完成"],
                        ["usa_today, washington_post", "128", "-", "-", "-", "CUDA OOM"],
                        ["usa_today, washington_post", "64", "0.8032", "0.8032", "0.8028", "完成"],
                    ],
                    "注：usa_today,washington_post 在 batch size 128 下受显存限制，后续分析采用 batch size 64 的完成结果。",
                    [2.15, 0.7, 0.75, 0.75, 0.75, 1.15],
                ),
                Table(
                    "表4  COVE-lite 构造与字段覆盖",
                    ["检查内容", "数值"],
                    [
                        ["扫描 JSON 文件", str(metrics["ctx"]["json_files"])],
                        ["可用样本数", str(metrics["ctx"]["available_before_cap"])],
                        ["本次保留样本", str(metrics["ctx"]["kept"])],
                        ["missing id/text/context", f"{metrics['ctx']['missing_ids']} / {metrics['ctx']['missing_text']} / {metrics['ctx']['missing_true_context']}"],
                        ["current entity/location/time", f"{event['field_presence_counts']['current_entities']} / {event['field_presence_counts']['current_locations']} / {event['field_presence_counts']['current_times']}"],
                        ["current event_type/relation", f"{event['field_presence_counts']['current_event_types']} / {event['field_presence_counts']['current_relations']}"],
                    ],
                    "注：字段覆盖率用于判断 true-context oracle 分析的可比范围，不代表最终在线推理输入。",
                    [3.0, 3.5],
                ),
            ],
            [],
        ),
        (
            "5 结果与分析",
            [
                "反事实训练规模直接影响 no-true-context 归因效果。MaxPerType 从 80 增至 1000 时，logistic regression head 的 Type Acc 从 0.2745 提升到 0.5275，Field Micro-F1 从 0.3564 提升到 0.5719。单字段反事实提供了可学习的字段冲突信号，但 1000 设置中的 location 样本只有 797 条，反映出字段可抽取样本不完全均衡。",
                "加入 different-event 类后，5way_1000 设置的 logistic regression 得到 Type Acc 0.4011 和 Field Micro-F1 0.5841。plus2000 设置将 different-event 训练样本增加到 3000 条，默认模型变为 image+caption MLP，合成测试集 Type Acc 为 0.5220，Field Micro-F1 为 0.6876。该变化与人工标注发现一致：真实 OOC 主要不是单字段替换，而是多字段事件冲突。",
            ],
            [
                Table(
                    "表5  no-true-context 反事实训练规模结果",
                    ["MaxPerType", "训练分布", "方法", "Type Acc", "Field F1", "Exact"],
                    [
                        ["80", "80/80/80/80/0", "LR", "0.2745", "0.3564", "0.1961"],
                        ["200", "200/200/200/200/0", "LR", "0.4266", "0.5195", "0.2308"],
                        ["1000", "1000/797/1000/1000/0", "LR", "0.5275", "0.5719", "0.3250"],
                    ],
                    "注：训练分布顺序为 none/location/time/entity/different-event；LR 表示 logistic regression no-true-context。",
                    [1.0, 1.75, 1.0, 0.95, 0.95, 0.85],
                ),
                Table(
                    "表6  five-class 与 plus2000 合成测试",
                    ["设置", "训练分布", "选中模型", "N", "Type Acc", "Field F1", "Exact"],
                    [
                        ["5way_1000", "1000/1000/1000/1000/987", "LR", "703", fmt(five["selected_model_score"]["mismatch_type_accuracy"]), fmt(five["selected_model_score"]["conflict_field_micro_f1"]), fmt(five["selected_model_score"]["exact_match_rate"])],
                        ["plus2000", "1000/1000/1000/1000/3000", "MLP", "998", fmt(plus["selected_model_score"]["mismatch_type_accuracy"]), fmt(plus["selected_model_score"]["conflict_field_micro_f1"]), fmt(plus["selected_model_score"]["exact_match_rate"])],
                    ],
                    "注：训练分布顺序为 none/entity/location/time/different-event；MLP 为 image+caption attribution head。",
                    [1.2, 2.25, 0.8, 0.65, 0.9, 0.9, 0.8],
                ),
            ],
            [
                Figure(
                    "图2  no-true-context 归因头随训练规模变化",
                    figs["scaling"],
                    "图2显示反事实训练样本增加后，Type Acc 和 Field Micro-F1 同时提高。该趋势只支持受控反事实分布下的学习效果，不推出真实 OOC 已被完全覆盖。",
                )
            ],
        ),
        (
            "6 真实 OOC 人工评估",
            [
                "人工评估集包含两批共 100 条真实 OOC 样本。标注对象为 current caption 与 true image context 的错配类型和冲突字段；标注时保留图文内容、原始上下文和候选字段，不使用模型预测作为主统计依据。本轮数据完成规范化单人标注，未形成双人独立盲审、Cohen's kappa 或第三人仲裁，因此人工集用于真实分布评估和误差定位，不用于显著性检验。",
                "类型分布中 different-event mismatch 为 85 条，entity mismatch 为 10 条，location mismatch 为 3 条，context omission 和 event-type mismatch 各 1 条。字段分布中 entity 为 94，event_type 为 74，relation 为 69，location 为 57，time 为 33。该分布说明真实 OOC 多数同时涉及多个字段。",
                "plus2000 模型在真实 OOC 上将 Type Acc 从 0.0900 提高到 0.2900，Field Micro-F1 从 0.3276 提高到 0.4781。Exact Match 仍为 0.0300，说明模型能恢复部分字段，但尚不能稳定给出完整事件冲突集合。",
            ],
            [
                Table(
                    "表7  真实 OOC 人工 100 条分布",
                    ["类别", "数量"],
                    [
                        ["different-event", "85"],
                        ["entity", "10"],
                        ["location", "3"],
                        ["context omission", "1"],
                        ["event-type", "1"],
                    ],
                    "注：字段分布为 entity 94、event_type 74、relation 69、location 57、time 33，存在多字段共现。",
                    [3.2, 1.5],
                ),
                Table(
                    "表8  真实 OOC 100 条 no-true-context 评估",
                    ["模型", "Type Acc", "Field F1", "Exact", "主要预测"],
                    [
                        ["5way_1000", "0.0900", "0.3276", "0.0300", "entity 40；temporal 28；benign 26"],
                        ["plus2000", "0.2900", "0.4781", "0.0300", "entity 44；different-event 33；temporal 14"],
                    ],
                    "注：Field F1 为 conflict field micro-F1；主要预测只列高频类别。",
                    [1.15, 0.9, 0.9, 0.8, 2.75],
                ),
            ],
            [
                Figure(
                    "图3  真实 OOC 人工标注类型分布",
                    figs["manual"],
                    "图3显示 different-event 在人工 100 条中占 85%。因此，单字段反事实只覆盖真实分布的一部分，训练集中需要包含多字段或跨事件错配样本。",
                ),
                Figure(
                    "图4  plus2000 在真实 OOC 评估中的变化",
                    figs["real"],
                    "图4显示 plus2000 设置提高了 Type Acc 和 Field Micro-F1。Exact Match 未同步提高，说明当前模型仍偏向预测部分字段，完整归因仍需更强事件表示和更多人工标注。",
                ),
            ],
        ),
        (
            "7 原型系统",
            [
                "实现部分采用 Gradio 单页界面承载推理流程。输入为图片和新闻文本，输出包括 OOC 判定、归因类型、冲突字段和证据状态。服务启动时自动选择可用端口并执行模型预热；缺少训练产物或 CLIP 依赖时返回 evidence insufficient，避免把缺失特征解释为高置信 benign。",
                "复现实验和轻量推理分离。运行界面只需要源码、Python 依赖和轻量模型产物；重新训练或复现 VDT 需要 VisualNews、NewsCLIPpings、VDT/BLIP-2 checkpoint 及本地路径配置。",
            ],
            [
                Table(
                    "表9  实现检查",
                    ["检查项", "结果"],
                    [
                        ["编译检查", "demo/app.py 与推理脚本通过"],
                        ["单元测试", "pytest -q：16 passed"],
                        ["结构检查", "check_project.py：OK"],
                        ["最终检查", "check_final_deliverables.py：OK"],
                        ["HTTP 访问", "Gradio 页面返回 200 OK"],
                    ],
                    None,
                    [2.6, 3.9],
                )
            ],
            [],
        ),
        (
            "8 讨论",
            [
                "第一，原始 OOC 不能全部标为 different-event。抽查样本中存在同一人物不同活动、同类体育比赛错配、同主题政治报道错配等 hard negative。本文只把严格筛选的低相似样本作为 weak different-event 训练样本，并用人工集评估真实泛化。",
                "第二，true-context oracle 与 no-true-context 推理需要分开解释。COVE-lite 能利用 VisualNews 原始上下文，因此适合分析上限和构造监督；在线系统只使用 image 与 current caption，指标较低符合输入信息减少后的任务难度。",
                "第三，真实 OOC 的 Exact Match 仍低。人工标注显示 different-event 常与 entity、event_type、relation、location、time 多字段共同出现。当前 prompt grounding 和图文相似度特征能捕捉部分字段，不足以稳定恢复完整事件结构。",
            ],
            [],
            [],
        ),
        (
            "9 结论",
            [
                "本文得到三点发现。第一，VDT strict baseline 可以在本地显存约束下完成核心复现，batch size 变化需要随结果一并记录。第二，可控反事实样本能为 no-true-context 归因头提供字段监督，训练规模增加后 Type Acc 和 Field Micro-F1 均提高。第三，真实 OOC 的主体、事件类型、关系、地点和时间常共同变化，plus2000 different-event 样本能改善部分字段预测，但完整 Exact Match 仍低。",
                "这一结果要求后续研究同时控制主分类、归因标签和真实事件结构。可执行方向包括扩大人工归因标注、引入双人盲审与第三人仲裁、构造 image_A + caption_B 的严格跨事件反事实样本、补充 OCR/NER/captioning 特征，并在不降低 VDT 主分类性能的条件下研究事件字段与主分类的受控融合。",
            ],
            [],
            [],
        ),
        (
            "10 作者贡献",
            [
                "费思岳参与数据路径整理和 VDT strict 复现配置核对，记录 NewsCLIPpings/VisualNews 数据状态、目标域设置和显存约束下的 batch size 调整。",
                "陈宇欣参与真实 OOC 样本检查、错配字段定义和人工标注表整理，支持 different-event 与多字段共现的案例分析。",
                "朱俊好负责研究路线、核心实现与实验集成，包括 VDT 接入、COVE-lite 构造、可控反事实训练、plus2000 评估和 no-true-context 推理系统。",
                "郭远重参与展示材料和复现说明整理，协助准备演示样例、答辩材料和轻量运行环境说明。",
            ],
            [],
            [],
        ),
        (
            "11 数据、代码与补充材料可用性声明",
            [
                "本研究代码、实验脚本、配置说明和轻量演示材料整理在课程项目仓库中。NewsCLIPpings、VisualNews 和 VDT/BLIP-2 checkpoint 按原数据集和模型许可获取，不随本文重新分发。人工评估表仅保留样本编号、字段标签和必要文本上下文，提交版本不包含额外个人信息。",
            ],
            [],
            [],
        ),
    ]

    refs = [
        "Luo G, Darrell T, Rohrbach A. NewsCLIPpings: Automatic Generation of Out-of-Context Multimodal Media. EMNLP, 2021.",
        "Liu F, Wang Y, Wang T, Ordonez V. Visual News: Benchmark and Challenges in News Image Captioning. EMNLP, 2021.",
        "Radford A, Kim J W, Hallacy C, et al. Learning Transferable Visual Models From Natural Language Supervision. ICML, 2021.",
        "Yang X, Zhang H, Lin Z, Hu Y, Han H. Out-of-Context Misinformation Detection via Variational Domain-Invariant Learning with Test-Time Training. AAAI, 2026.",
        "Qi P, Yan Z, Hsu W, Lee M L. SNIFFER: Multimodal Large Language Model for Explainable Out-of-Context Misinformation Detection. CVPR, 2024.",
        "Tonglet J, Thiem G, Gurevych I. COVE: COntext and VEracity prediction for out-of-context images. NAACL, 2025.",
        "Hugging Face. Transformers Documentation: CLIPModel and BART MNLI APIs.",
        "Gradio Team. Gradio Documentation: Blocks, JSON, HTML components and queue configuration.",
    ]

    return abstract, sections, refs


def write_markdown(path: Path, abstract: list[str], sections, refs):
    lines = [f"# {TITLE}", "", EN_TITLE, "", "作者：费思岳，陈宇欣，朱俊好，郭远重", "", "## 摘要"]
    lines.extend(abstract)
    lines.append("")
    lines.append("关键词：内容安全；图文内容挪用；Out-of-Context misinformation；VDT；可控反事实；错配归因")
    for title, paras, tables, figs in sections:
        lines += ["", f"## {title}"]
        for p in paras:
            lines.append("")
            lines.append(p)
        for table in tables:
            lines += ["", table.caption, "", "|" + "|".join(table.header) + "|", "|" + "|".join(["---"] * len(table.header)) + "|"]
            for row in table.rows:
                lines.append("|" + "|".join(row) + "|")
            if table.note:
                lines += ["", table.note]
        for fig in figs:
            lines += ["", fig.caption, "", f"![{fig.caption}](paper_assets/{fig.path.name})", "", fig.note]
    lines += ["", "## 参考文献"]
    for idx, ref in enumerate(refs, 1):
        lines.append(f"[{idx}] {ref}")
    path.write_text("\n".join(lines), encoding="utf-8")


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=False, color="000000"):
    run.font.name = west
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)


def set_para(p, first=True, align=None, line=1.5, before=0, after=0):
    p.paragraph_format.first_line_indent = Pt(21) if first else Pt(0)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc: Document, text: str, first=True, align=None, size=10.5, bold=False, east="宋体"):
    p = doc.add_paragraph()
    set_para(p, first=first, align=align)
    r = p.add_run(text)
    set_run_font(r, east=east, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_para(p, first=False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # clear auto numbering, add explicit text
    r = p.add_run(text)
    set_run_font(r, east="黑体", size=14 if level == 1 else 12, bold=True)
    return p


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for k, v in attrs.items():
            element.set(qn(f"w:{k}"), str(v))


def set_cell_margins(cell, top=60, bottom=60, start=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float] | None):
    if not widths:
        return
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_table(doc: Document, table_data: Table):
    p = doc.add_paragraph()
    set_para(p, first=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(table_data.caption)
    set_run_font(r, size=9)
    rows = [table_data.header] + table_data.rows
    table = doc.add_table(rows=len(rows), cols=len(table_data.header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, table_data.widths)
    for i, row in enumerate(rows):
        tr = table.rows[i]
        if i == 0:
            tr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for j, value in enumerate(row):
            cell = tr.cells[j]
            clear_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_para(p, first=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
            r = p.add_run(value)
            set_run_font(r, east="黑体" if i == 0 else "宋体", size=9.5 if i == 0 else 9, bold=(i == 0))
            if i == 0:
                set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"}, bottom={"val": "single", "sz": "6", "color": "000000"})
            if i == len(rows) - 1:
                set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})
    if table_data.note:
        note = doc.add_paragraph()
        set_para(note, first=False, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)
        rn = note.add_run(table_data.note)
        set_run_font(rn, size=9)
    doc.add_paragraph()


def add_figure(doc: Document, fig: Figure):
    p = doc.add_paragraph()
    set_para(p, first=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    p.add_run().add_picture(str(fig.path), width=Inches(5.85))
    cap = doc.add_paragraph()
    set_para(cap, first=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run(fig.caption)
    set_run_font(r, size=9)
    add_para(doc, fig.note, first=True)


def add_toc(doc: Document):
    p = doc.add_paragraph()
    set_para(p, first=False)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string("000000")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)


def build_docx(path: Path, abstract: list[str], sections, refs):
    doc = Document()
    configure_doc(doc)
    add_para(doc, TITLE, first=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, east="黑体")
    add_para(doc, EN_TITLE, first=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    add_para(doc, "费思岳，陈宇欣，朱俊好，郭远重", first=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    add_para(doc, "武汉大学国家网络安全学院", first=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    add_para(doc, "课程：内容安全；答辩序号：10", first=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    doc.add_paragraph()
    add_para(doc, "摘要", first=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, bold=True, east="黑体")
    for p in abstract:
        add_para(doc, p)
    add_para(doc, "关键词：内容安全；图文内容挪用；Out-of-Context misinformation；VDT；可控反事实；错配归因", first=False)
    for title, paras, tables, figs in sections:
        add_heading(doc, title, 1)
        for p in paras:
            add_para(doc, p)
        for table in tables:
            add_table(doc, table)
        for fig in figs:
            add_figure(doc, fig)
    add_heading(doc, "参考文献", 1)
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        set_para(p, first=False, line=1.0)
        r = p.add_run(f"[{idx}] {ref}")
        set_run_font(r, size=9.5)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "内容安全课程论文"
    doc.core_properties.author = "费思岳；陈宇欣；朱俊好；郭远重"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = {
        "plus": load_json("outputs/no_true_context_attr_5way_plus2000/no_true_context_attr_metrics.json"),
        "five": load_json("outputs/no_true_context_attr_5way_1000/no_true_context_attr_metrics.json"),
        "manual": load_json("examples/real_ooc_manual_100_summary.json"),
        "real_eval": load_json("outputs/real_ooc_no_true_context_eval_metrics.json"),
        "ctx": load_json("outputs/cove_lite_context_pairs_3000.jsonl.stats.json"),
        "event": load_json("outputs/event_tuples_v2.jsonl.stats.json"),
    }
    figs = build_assets(metrics)
    repo_asset_dir = REPO / "docs" / "report" / "paper_assets"
    repo_asset_dir.mkdir(parents=True, exist_ok=True)
    for fig_path in figs.values():
        shutil.copy2(fig_path, repo_asset_dir / fig_path.name)
    abstract, sections, refs = content(metrics, figs)
    md = REPORT_DIR / "10-VDT-CF-Attr跨域图文内容挪用检测与错配归因-专业论文版.md"
    docx = REPORT_DIR / "10-VDT-CF-Attr跨域图文内容挪用检测与错配归因-专业论文版.docx"
    repo_md = REPO / "docs" / "report" / "FINAL_COURSE_PAPER_PROFESSIONAL.md"
    write_markdown(md, abstract, sections, refs)
    repo_md.write_text(md.read_text(encoding="utf-8"), encoding="utf-8")
    build_docx(docx, abstract, sections, refs)
    print(docx)
    print(md)
    print(repo_md)


if __name__ == "__main__":
    main()

